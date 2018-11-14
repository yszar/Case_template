import os
import json
import pathlib
import docxtpl
import pinyin
import time
import glob
import re
import docx
path = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))

# path = os.path.dirname(__file__)


def load_old_json(info):
    # module_path = os.path.dirname(__file__)
    with open(path + '/' + info + '.json', 'r') as f:
        data = json.load(f)
    return data


def replacedocx(save_path, docxname, new_json):
    # module_path = os.path.dirname(__file__)
    doc = docxtpl.DocxTemplate(path + "/" + "package/" + 'docxs' + "/" +
                               docxname + ".docx")
    doc.render(new_json)
    doc.save(save_path + "/" + docxname + '.docx')


def get_dict_value(date, keys, default=None):
    # default=None，在key值不存在的情况下，返回None
    keys_list = keys.split('.')
    # 以“.”为间隔，将字符串分裂为多个字符串，其实字符串为字典的键，保存在列表keys_list里
    if isinstance(date, dict):
        # 如果传入的数据为字典
        dictionary = dict(date)
        # 初始化字典
        for i in keys_list:
            # 按照keys_list顺序循环键值
            try:
                # if dictionary.get(i) != None:
                if dictionary.get(i) is not None:
                    dict_values = dictionary.get(i)
                # 如果键对应的值不为空，返回对应的值
                # elif dictionary.get(i) == None:
                elif dictionary.get(i) is None:
                    dict_values = dictionary.get(int(i))
                # 如果键对应的值为空，将字符串型的键转换为整数型，返回对应的值
            except Exception:
                return default
                # 如果字符串型的键转换整数型错误，返回None
            dictionary = dict_values
        return dictionary
    else:
        # 如果传入的数据为非字典
        try:
            dictionary = dict(eval(date))
            # 如果传入的字符串数据格式为字典格式，转字典类型，不然返回None
            if isinstance(dictionary, dict):
                for i in keys_list:
                    # 按照keys_list顺序循环键值
                    try:
                        if dictionary.get(i) is not None:
                            dict_values = dictionary.get(i)
                        # 如果键对应的值不为空，返回对应的值
                        elif dictionary.get(i) is None:
                            dict_values = dictionary.get(int(i))
                        # 如果键对应的值为空，将字符串型的键转换为整数型，返回对应的值
                    except Exception:
                        return default
                        # 如果字符串型的键转换整数型错误，返回None
                    dictionary = dict_values
                return dictionary
        except Exception:
            return default


def mkdir_save_path(old_json):
    # old_json = load_old_json()
    if '超过食品安全标准限量的' in old_json['illegal_label']:
        save_path = (path + '/' + old_json['company_name'] + old_json['har'] +
                     old_json['illegal_label'] + old_json['food'] + '案')
    # return save_path
    elif '超过食品安全标准限量的' not in old_json['illegal_label']:
        save_path = (path + '/' + old_json['company_name'] +
                     old_json['illegal_label'] + '案')
    pathlib.Path(save_path).mkdir(parents=True, exist_ok=True)
    return save_path


def lawopen(txt):
    # path = package.docxreplace.path
    with open(path + txt, 'r') as f:
        data = json.load(f)
    return data


def law_json(name, old_json):
    name = pinyin.get(name, format="strip", delimiter=" ")
    law = lawopen("/law/" + name + ".json")

    old_json['violation'] = get_dict_value(
        law, old_json['illegal_label'] + ".violation")
    old_json['violation_content'] = get_dict_value(
        law, old_json['illegal_label'] + ".violation_content")
    old_json['according'] = get_dict_value(
        law, old_json['illegal_label'] + ".according")
    old_json['according_content'] = get_dict_value(
        law, old_json['illegal_label'] + ".according_content")
    return old_json


def age(number):
    current = int(time.strftime("%Y"))
    year = int(number[6:10])
    age = current - year
    return age


def changejson(new_json, *the_object):
    # json = package.docxreplace.package.docxreplace()

    if the_object == ('2',):
        print('输入授权委托人职务')
        new_json['temp_position'] = input()
        print('输入授权委托人姓名')
        new_json['templr'] = input()
        print('输入授权委托人身份证号码')
        new_json['tempid'] = input()

        new_json['legal_representative'], new_json['templr'] = new_json[
            'templr'], new_json['legal_representative']
        new_json['identification_number'], new_json['tempid'] = new_json[
            'tempid'], new_json['identification_number']
        new_json['position'], new_json['temp_position'] = new_json[
            'temp_position'], new_json['position']

    finally_id = new_json['identification_number']
    if int(finally_id[-2]) % 2 == 0:
        new_json['gender'] = '女'
    else:
        new_json['gender'] = '男'
    new_json['age'] = str(age(new_json['identification_number']))

    if new_json['category'] == '食':
        new_json['business'] = '食品经营'
        new_json['law_name'] = '《中华人民共和国食品安全法》'
        new_json = law_json(new_json['category'], new_json)
        if '超过食品安全标准限量的' in new_json['illegal_label']:
            new_json['illegal_behavior'] = (
                new_json['har'] + new_json['illegal_behavior'] +
                new_json['food'])
        else:
            new_json['illegal_behavior'] = new_json['illegal_label']
    elif new_json['category'] == '药':
        new_json['business'] = '药品经营'
    elif new_json['category'] == '妆':
        new_json['business'] = '化妆品经营'
    elif new_json['category'] == '械':
        new_json['business'] = '医疗器械经营'
    # if new_json['illegal_behavior'] == '超过食品安全标准限量的':
    #     print('输入超限食品名称')
    #     food_name = input()
    #     print('输入超限的元素')
    #     harmful = input()
    #     new_json['fullib'] = harmful + '超过食品安全标准限量的' + food_name

    return new_json


def all_json_name():
    filenames = glob.glob(path + r'/*.json')
    patten = 'Case_template/(.*).json'
    allname = []
    for i in filenames:
        jsonname = re.findall(patten, i)
        allname.append(jsonname[0])
    print(allname)


def get_happening(path):
    '''获取文件'''
    # 获得word文档
    docxfile = docx.Document(path)
    # print(file)
#     return docxfile


# def preproccess_file(docxfile):
#     '''文件预处理'''
    # 输出文档段落数（行数）
    # paragraph_sum = len(file.paragraphs)
    # print(paragraph_sum)

    # 输出每一段的内容
    para_list = []
    for para in docxfile.paragraphs:
        # print(para.text)
        para_list.append(para.text)

    # 合并字符串
    file_text = ''.join(para_list)
    # print(file_text)
    # return file_text


# def extract_file(file_text):
#     '''提取内容'''
    # 使用正则提取关键字后面的数字
    result = re.findall('基本情况介绍：(.*)附件：', file_text)
    # print(result)
    return result


# old_json = docxreplace()

# save_path = path