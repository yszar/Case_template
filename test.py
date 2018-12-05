from flask import Flask, request, render_template, send_file
from docx import Document
import io
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import datetime
from flask_wtf import FlaskForm
from wtforms import StringField, SubmitField, SelectField
from wtforms.validators import Required
from flask_bootstrap import Bootstrap
from urllib.parse import quote

app = Flask(__name__)
app.config['SECRET_KEY'] = '111111'

bootstrap = Bootstrap(app)


class NameForm(FlaskForm):
    name = StringField('姓名')
    status = SelectField(
        '按类型查询',
        validators=[Required()],
        choices=[('0', '全部'), ('1', '待审核'), ('2', '认证成功'), ('3', '认证失败')])
    status1 = SelectField(
        '按类型查询', validators=[Required()], choices=[
            ('0', '全部'),
        ])
    departments = StringField('院系')
    specialty = StringField('专业')
    class_name = StringField('班级：')
    gender = StringField('性别：')
    school_num = StringField('学号：')
    school_time_year = StringField('在校起始时间年')
    school_time_month = StringField('在校起始时间月')
    school_time_year1 = StringField('在校截止时间年')
    school_time_month1 = StringField('在校截止时间月')
    home_address = StringField('家庭通讯地址：')
    personal_tel = StringField('个人联系方式：')
    home_num = StringField('家庭联系方式：')
    reason = StringField('本人申请理由：')
    submit = SubmitField('申请表下载')
    submit1 = SubmitField('清单下载')


@app.route('/', methods=['GET', 'POST'])
def index():
    name = None
    form = NameForm()
    if form.validate_on_submit():
        # if request.form['stname'] == 'submit':
        # name = NameForm.data['name']
        name = form.name.data
        departments = form.departments.data
        specialty = form.specialty.data
        class_name = form.class_name.data
        gender = form.gender.data
        school_num = form.school_num.data
        school_time_year = form.school_time_year.data
        school_time_month = form.school_time_month.data
        school_time_year1 = form.school_time_year1.data
        school_time_month1 = form.school_time_month1.data
        home_address = form.home_address.data
        personal_tel = form.personal_tel.data
        home_num = form.home_num.data
        reason = form.reason.data

        form.name.data = ''
        form.departments.data = ''
        form.specialty.data = ''
        form.class_name.data = ''
        form.gender.data = ''
        form.school_num.data = ''
        form.school_time_year.data = ''
        form.school_time_month.data = ''
        form.school_time_year1.data = ''
        form.school_time_month1.data = ''
        form.home_address.data = ''
        form.personal_tel.data = ''
        form.home_num.data = ''
        form.reason.data = ''

        nt = datetime.datetime.now()

        if request.form['key'] == '申请表下载':
            document = Document()
            document.styles['Normal'].font.name = u'宋体'
            document.styles['Normal']._element.rPr.rFonts.set(
                qn('w:eastAsia'), u'宋体')
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('X X X X X X 学 院\rXXX申请表')
            r.font.size = Pt(16)
            r.bold = True

            table = document.add_table(rows=37, cols=13, style='Table Grid')
            table.autofit = False
            table.columns[0].width = Inches(0.49)
            table.cell(0, 0).merge(table.cell(2, 2))
            table.cell(0, 3).merge(table.cell(2, 6))
            table.cell(0, 7).merge(table.cell(2, 9))
            table.cell(0, 10).merge(table.cell(2, 12))
            table.cell(3, 0).merge(table.cell(5, 2))
            table.cell(3, 3).merge(table.cell(5, 5))
            table.cell(3, 6).merge(table.cell(5, 6))
            table.cell(3, 7).merge(table.cell(5, 9))
            table.cell(3, 10).merge(table.cell(5, 10))
            table.cell(3, 11).merge(table.cell(5, 12))
            table.cell(6, 0).merge(table.cell(8, 2))
            table.cell(6, 3).merge(table.cell(8, 12))
            table.cell(9, 0).merge(table.cell(14, 2))
            table.cell(9, 3).merge(table.cell(14, 6))
            table.cell(9, 7).merge(table.cell(11, 9))
            table.cell(12, 7).merge(table.cell(14, 9))
            table.cell(9, 10).merge(table.cell(11, 12))
            table.cell(12, 10).merge(table.cell(14, 12))
            table.cell(15, 0).merge(table.cell(20, 2))
            table.cell(15, 3).merge(table.cell(20, 12))
            table.cell(21, 0).merge(table.cell(26, 2))
            table.cell(21, 3).merge(table.cell(26, 12))
            table.cell(27, 0).merge(table.cell(31, 2))
            table.cell(27, 3).merge(table.cell(31, 12))
            table.cell(32, 0).merge(table.cell(36, 2))
            table.cell(32, 3).merge(table.cell(36, 12))
            hdr_cells0 = table.rows[0].cells
            hdr_cells3 = table.rows[3].cells
            hdr_cells6 = table.rows[6].cells
            hdr_cells9 = table.rows[9].cells
            hdr_cells12 = table.rows[12].cells
            hdr_cells15 = table.rows[15].cells
            hdr_cells21 = table.rows[21].cells
            hdr_cells27 = table.rows[27].cells
            hdr_cells32 = table.rows[32].cells

            hdr_cells0[0].add_paragraph(
                '院（系）\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells0[3].add_paragraph(
                departments + '院').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells0[7].add_paragraph(
                '专业班级').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells0[10].add_paragraph(
                specialty + class_name).alignment = WD_ALIGN_PARAGRAPH.CENTER

            hdr_cells3[0].add_paragraph(
                '姓名\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells3[3].add_paragraph(
                name).alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells3[6].add_paragraph(
                '性别').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells3[7].add_paragraph(
                gender).alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells3[10].add_paragraph(
                '学号').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells3[11].add_paragraph(
                school_num).alignment = WD_ALIGN_PARAGRAPH.CENTER

            hdr_cells6[0].add_paragraph(
                '在校时间\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells6[3].add_paragraph(
                school_time_year + '年' + school_time_month + '月——' +
                school_time_year1 + '年' + school_time_month1 +
                '月').alignment = WD_ALIGN_PARAGRAPH.CENTER

            hdr_cells9[0].add_paragraph(
                '家庭通讯地址').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells9[3].add_paragraph(
                home_address).alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells9[7].add_paragraph(
                '家庭联系方式').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells12[7].add_paragraph(
                '个人联系方式').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells9[10].add_paragraph(
                home_num).alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells12[10].add_paragraph(
                personal_tel).alignment = WD_ALIGN_PARAGRAPH.CENTER

            hdr_cells15[0].add_paragraph(
                '\n\n本人申请理由').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells15[3].add_paragraph(
                reason + '\n\n学生本人签名：' +
                '\n\n\t\t年\t月\t日').alignment = WD_ALIGN_PARAGRAPH.CENTER

            hdr_cells21[0].add_paragraph(
                '\n\n院系领导意见').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells21[3].add_paragraph(
                '\n\n领导签字：' +
                '\n\n\t\t年\t月\t日').alignment = WD_ALIGN_PARAGRAPH.CENTER

            hdr_cells27[0].add_paragraph(
                '\n\n学生处意见').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells27[3].add_paragraph(
                '\n\n领导签字：' +
                '\n\n\t\t年\t月\t日').alignment = WD_ALIGN_PARAGRAPH.CENTER

            hdr_cells32[0].add_paragraph(
                '\n\n主管校领导审批').alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells32[3].add_paragraph(
                '\n\n领导签字：' +
                '\n\n\t\t年\t月\t日').alignment = WD_ALIGN_PARAGRAPH.CENTER

            # document.save(name + 'XX申请表.docx')
            f = io.BytesIO()
            document.save(f)
            length = f.tell()
            f.seek(0)
            filename = quote(name + '申请表.doc')
            rv = send_file(f, as_attachment=True, attachment_filename=filename)
            rv.headers[
                'Content-Disposition'] += "; filename*=utf-8''{}".format(
                    filename)
            return rv
        elif request.form['key'] == '清单下载':
            document1 = Document()
            document1.styles['Normal'].font.name = u'宋体'
            document1.styles['Normal']._element.rPr.rFonts.set(
                qn('w:eastAsia'), u'宋体')
            p1 = document1.add_paragraph()
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.line_spacing = Pt(25)
            r1 = p1.add_run('X X X X X X X X\rX X 清 单\n')
            r1.font.size = Pt(16)
            r1.bold = True

            leave_name = document1.add_paragraph('各部门：\n\n\t')
            leave_name.add_run(departments).font.underline = True
            leave_name.add_run('院（系）')
            leave_name.add_run(specialty).font.underline = True
            leave_name.add_run('专业')
            leave_name.add_run(class_name).font.underline = True
            leave_name.add_run('班,学生')
            leave_name.add_run(name).font.underline = True
            leave_name.add_run('，（学号：')
            leave_name.add_run(school_num).font.underline = True
            leave_name.add_run('）。申请')
            leave_name.add_run('XX').font.underline = True
            leave_name.add_run(',请予以协助办理相关XX手续。')
            leave_name.add_run(
                '\n\n\n\n\t\t\t\t\t\t\t\t\tX X 处  （公 章）\n\n\t\t\t\t\t\t\t\t\t')
            leave_name.add_run(
                nt.strftime('%Y{y}%m{m}%d{d}\n\n\n\n\n\n\n').format(
                    y='年', m='月', d='日'))

            table1 = document1.add_table(rows=10, cols=8, style='Table Grid')
            table1.cell(0, 0).merge(table1.cell(7, 3))
            table1.cell(0, 4).merge(table1.cell(7, 7))
            table1.cell(8, 0).merge(table1.cell(9, 7))

            hdr1_cells0 = table1.rows[0].cells
            hdr1_cells1 = table1.rows[8].cells

            dormitory = hdr1_cells0[0].add_paragraph()
            dormitory.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dormitory1 = dormitory.add_run('XXX\n\n\n（盖章）\n\n\t\t年  月  日')
            dormitory1.bold = True
            dormitory1.font.size = Pt(14)

            library = hdr1_cells0[4].add_paragraph()
            library.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            library1 = library.add_run('XXX\n\n\n（盖章）\n\n\t\t年  月  日')
            library1.bold = True
            library1.font.size = Pt(14)

            note = hdr1_cells1[0].add_paragraph()
            note1 = note.add_run('备注：\n')
            note1.bold = True
            note1.font.size = Pt(12)

            f = io.BytesIO()
            document1.save(f)
            length = f.tell()
            f.seek(0)
            filename = quote(name + '清单.doc')
            rv = send_file(f, as_attachment=True, attachment_filename=filename)
            rv.headers[
                'Content-Disposition'] += "; filename*=utf-8''{}".format(
                    filename)
            return rv

    return render_template('index.html', form=form, name=name)


if __name__ == '__main__':
    app.run()
